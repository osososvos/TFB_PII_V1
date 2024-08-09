import re
import tkinter as tk
from tkinter import filedialog, messagebox

import docx
import spacy
import unicodedata

# Load model
nlp = spacy.load("en_core_web_md")


# Check if the file is already open. Load and Read Word Document
def load_document(file_path):
    try:
        return docx.Document(file_path)
    except docx.opc.exceptions.PackageNotFoundError:
        messagebox.showerror("Error", "El archivo está en uso. Por favor, ciérrelo antes de procesarlo.")
        return None


# Remove accents
def remove_accents(input_str):
    nfkd_form = unicodedata.normalize('NFKD', input_str)
    return "".join([c for c in nfkd_form if not unicodedata.combining(c)])


# Identify PII using NER and REGEX
def identify_pii(text):
    normalized_text = remove_accents(text)
    doc = nlp(normalized_text)
    pii_entities = []

    # Use SpaCy for standard PII detection
    for ent in doc.ents:
        if ent.label_ in ["PERSON", "GPE", "ORG", "DATE", "TIME", "ORDINAL", "CARDINAL", "LOC"]:
            pii_entities.append((ent.start_char, ent.end_char, ent.label_))

    # Custom regex for email addresses
    email_regex = re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
    for match in email_regex.finditer(text):
        pii_entities.append((match.start(), match.end(), "EMAIL"))

    # Custom regex to target alphanumeric patterns preceded by "number:" word.
    id_regex = re.compile(r'\bnumber:\s*([a-zA-Z0-9]+)\b', re.IGNORECASE)
    for match in id_regex.finditer(text):
        start, end = match.span(1)
        pii_entities.append((start, end, "ID"))

    print(pii_entities)
    return pii_entities


#  Redact PII with the entity name
def redact_text(text, pii_entities):
    redacted_text = text
    offset = 0
    for start, end, label in pii_entities:
        replacement_text = f"[{label}]"
        redacted_text = redacted_text[:start + offset] + replacement_text + redacted_text[end + offset:]
        offset += len(replacement_text) - (end - start)
    return redacted_text


#  Check if a paragraph contains a page break.
def has_page_break(para):
    return para._element.xpath('.//w:br[@w:type="page"]')


#  Inserts a page break in the document after the specified paragraph
def add_page_break(para):
    run = para.add_run()
    run.add_break(docx.text.run.WD_BREAK.PAGE)


# Process Text, Paragraphs, Breaks and Replace entities ###
def redact_document(doc):
    for para in doc.paragraphs:
        # Preserve paragraphs with page breaks
        if has_page_break(para):
            continue  # The loop skips this paragraph (to avoid altering it and to preserve the page break).
        text = para.text
        pii_entities = identify_pii(text)
        redacted_text = redact_text(text, pii_entities)
        para.text = redacted_text
    return doc


#  Save Redacted Document
def save_document(doc, output_path):
    doc.save(output_path)


# Main function to tie everything
def main(input_file, output_file):
    doc = load_document(input_file)
    redacted_doc = redact_document(doc)
    save_document(redacted_doc, output_file)


#  GUI:
def browse_file():
    filename = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if filename:
        entry_file_path.delete(0, tk.END)
        entry_file_path.insert(0, filename)


def redact_file():
    input_doc = entry_file_path.get()
    if not input_doc:
        messagebox.showwarning("Advertencia", "Por favor, seleccione un archivo .docx")
        return

    output_doc = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if not output_doc:
        return

    main(input_doc, output_doc)
    messagebox.showinfo("Información", "El archivo ha sido anonimizado y guardado correctamente.")


app = tk.Tk()
app.title("Eliminación de Información Personal Identificable en Documentos Word V1")

frame = tk.Frame(app, padx=10, pady=10)
frame.pack(padx=10, pady=10)

label_file_path = tk.Label(frame, text="Seleccionar archivo .docx:")
label_file_path.grid(row=0, column=0, sticky=tk.W)

entry_file_path = tk.Entry(frame, width=50)
entry_file_path.grid(row=1, column=0, pady=5)

button_browse = tk.Button(frame, text="Buscar", command=browse_file)
button_browse.grid(row=1, column=1, padx=5)

button_redact = tk.Button(frame, text="Anonimizar", command=redact_file)
button_redact.grid(row=2, column=0, columnspan=2, pady=10)

app.mainloop()
